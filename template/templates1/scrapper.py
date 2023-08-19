import requests
from bs4 import BeautifulSoup
import os
from docx import Document

# URL of the webpage to scrape
url = "https://www.flowrite.com/blog/thank-you-email"

# Fetch the webpage content
response = requests.get(url)
webpage_content = response.text

# Parse the HTML content using BeautifulSoup
soup = BeautifulSoup(webpage_content, "html.parser")

# Find all the divs with class "div-block-138"
email_divs = soup.find_all("div", class_="div-block-138")

# Create a directory to store the .docx files
output_directory = "email_templates"
if not os.path.exists(output_directory):
    os.makedirs(output_directory)

# Extract and create .docx files for the email templates
for email_div in email_divs:
    # Find the title of the email from the preceding h3 tag
    email_title = email_div.find_previous("h3")
    if email_title:
        template_title = ((email_title.get_text().strip()).split(".")[1]).strip()

        file_name = os.path.join(output_directory, f"{template_title}.docx")

        email_content = email_div.find("div", class_="paragraph-60")
        if email_content:
            email_text = email_content.get_text().strip()

            # Create a new document
            doc = Document()
            doc.add_heading(template_title, level=1)
            doc.add_paragraph(email_text)

            # Save the document as a .docx file
            doc.save(file_name)
            print(f"Created {file_name}")

print("Files created successfully.")
