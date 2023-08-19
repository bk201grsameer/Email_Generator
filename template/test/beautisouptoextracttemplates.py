from docx import Document
import os
from termcolor import cprint
import requests
from bs4 import BeautifulSoup

# Create a directory to store the .docx files
output_directory = "email_templates"
if not os.path.exists(output_directory):
    os.makedirs(output_directory)


def func():
    try:
        website = "https://www.flowrite.com/blog/thank-you-email"
        response = requests.get(website)
        if response.status_code == 200:
            soup = BeautifulSoup(response.content, "html.parser")

            # Find all the titles (assuming they are within <h3> tags)
            titles = soup.find_all("h3")
            # Find all the containers using CSS selector
            containers = soup.select("div.paragraph-60")

            if containers:
                for title, container in zip(titles, containers):
                    title_text = ((title.get_text(strip=True)).split("."))[1].strip()
                    file_name = os.path.join(output_directory, f"{title_text}.docx")
                    template_string = ""
                    for string in container.stripped_strings:
                        # print(string)
                        template_string += string + "\n"
                    createDocument(template_string, file_name)
                    cprint(f"[+] Created {file_name}", "green")
            else:
                print("Container not found.")

    except Exception as ex:
        print("[-] Exception:", str(ex))


def createDocument(email_text, file_name):
    try:
        doc = Document()
        doc.add_paragraph(email_text)
        # Save the document as a .docx file
        doc.save(file_name)
    except Exception as ex:
        print("[-]", str(ex))


if __name__ == "__main__":
    func()
